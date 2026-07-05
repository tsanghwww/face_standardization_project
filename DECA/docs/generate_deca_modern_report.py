from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/houwingtsang/Documents/face_standardization_project/DECA/docs/DECA项目现代化运行说明.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 45)
MUTED = RGBColor(91, 103, 112)
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def paragraph_border_bottom(paragraph, color="2E74B5", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def set_run(run, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        run = p.add_run(text)
        set_run(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text), size=10.5, color=INK)
    return p


def add_step(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text), size=10.5, color=INK)
    return p


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_table_borders(table, color="D7DEE8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_run(run, size=9.5, color=RGBColor(35, 45, 55), font="Courier New")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [2100, 7260])
    set_table_borders(table)
    for idx, row in enumerate(rows):
        if idx > 0:
            table.add_row()
        label, value = row
        cells = table.rows[idx].cells
        set_cell_shading(cells[0], LIGHT_GRAY)
        for c in cells:
            set_cell_margins(c)
        p0 = cells[0].paragraphs[0]
        set_run(p0.add_run(label), size=10.5, color=DARK_BLUE, bold=True)
        p1 = cells[1].paragraphs[0]
        set_run(p1.add_run(value), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(h), size=10, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx else WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(value), size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("DECA Modern Runtime Brief"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("Generated for local project handoff"), size=9, color=MUTED)

    add_para(doc, "技术交接文档", size=10, color=MUTED, bold=True, after=2)
    title = add_para(doc, "DECA 项目结构梳理与现代化运行说明", size=22, color=INK, bold=True, after=4)
    subtitle = add_para(
        doc,
        "范围：依赖环境升级、代码兼容改动、Mac 渲染路径、MPS 测试结果与运行建议",
        size=12,
        color=MUTED,
        after=10,
    )
    meta = add_para(doc, f"项目路径：/Users/houwingtsang/Documents/face_standardization_project/DECA  |  日期：{date.today().isoformat()}", size=10, color=MUTED, after=10)
    paragraph_border_bottom(meta, color="2E74B5", size="8")

    add_heading(doc, "执行摘要", 1)
    add_para(
        doc,
        "本次整理的目标是让原本面向 Python 3.7、PyTorch 1.6、CUDA 10.1 的 DECA 项目，在当前 Apple Silicon macOS 环境中使用现代 Python 与 PyTorch 跑通。项目已完成两类验证：无渲染模式可稳定生成关键点与 MATLAB 参数文件；安装 PyTorch3D 后，Mac 上也已成功生成可视化图、深度图、OBJ 与细节 OBJ。",
    )
    add_kv_table(
        doc,
        [
            ("当前结论", "现代环境已创建；无渲染、CPU 渲染、MPS 混合渲染均已跑通。"),
            ("本机环境", "macOS arm64；CUDA 不可用；MPS 可用；PyTorch3D 0.7.9 已源码编译安装。"),
            ("推荐路径", "批量运行优先用 CPU + PyTorch3D；MPS 可用但当前单图速度更慢。"),
            ("验证输出", "kpt2d/kpt3d、.mat、vis.jpg、depth.jpg、obj、detail.obj、texture/normal map 已生成。"),
        ],
    )

    add_heading(doc, "项目结构梳理", 1)
    add_para(doc, "项目主体位于 DECA 目录，是 SIGGRAPH 2021 DECA 的 PyTorch 实现，主要用于从单张人脸图像重建 3D 头部模型、表达参数与细节几何。")
    add_matrix(
        doc,
        ["目录/文件", "作用"],
        [
            ("demos/", "运行入口：demo_reconstruct.py、demo_transfer.py、demo_teaser.py。"),
            ("decalib/", "核心库：模型、FLAME、渲染器、数据集、工具函数。"),
            ("decalib/models/", "ResNet encoder、FLAME、decoder、lbs 等模型组件。"),
            ("decalib/utils/", "渲染、旋转转换、可视化、OBJ 输出、训练辅助。"),
            ("data/", "模型权重、FLAME 模型、纹理、UV mask 等数据资产。"),
            ("TestSamples/", "示例输入图、已有结果和测试素材。"),
            ("configs/", "训练配置文件。"),
            ("requirements.txt", "原始旧依赖：Python 3.7、PyTorch 1.6 路线。"),
            ("requirements_modern.txt", "新增现代依赖入口。"),
            ("RUNNING_MODERN.md", "新增中文现代运行说明。"),
        ],
        [2100, 7260],
    )

    add_heading(doc, "依赖与环境现状", 1)
    add_para(doc, "原始依赖与当前机器存在明显年代差异：项目默认 CUDA，且 standard rasterizer 会编译 .cu 扩展；这在 Apple Silicon macOS 上不能直接使用。")
    add_matrix(
        doc,
        ["项目/组件", "原始要求", "本次现代化结果"],
        [
            ("Python", "3.7", "conda 环境 deca-modern：Python 3.10.20"),
            ("PyTorch", "1.6.0", "torch 2.12.0"),
            ("torchvision", "0.7.0", "torchvision 0.27.0"),
            ("torchaudio", "未要求", "torchaudio 2.11.0"),
            ("CUDA", "默认 cuda", "本机 CUDA False；MPS True"),
            ("chumpy", "旧包，依赖旧 NumPy API", "通过 no-build-isolation 安装，并在代码中兼容旧别名"),
            ("renderer", "standard CUDA 或 pytorch3d", "PyTorch3D 0.7.9 已在 macOS arm64 源码编译成功；standard CUDA 路线仍不适用于 Mac"),
        ],
        [1700, 2500, 5160],
    )

    add_heading(doc, "本次代码与文件变动", 1)
    add_para(doc, "改动保持在兼容运行与说明文档范围内，没有重写算法主体。")
    add_matrix(
        doc,
        ["文件", "变动说明"],
        [
            ("requirements_modern.txt", "新增现代依赖清单，不直接包含 torch，方便按平台选择 PyTorch wheel。"),
            ("environment_modern.yml", "新增 conda 环境定义，使用 Python 3.10。"),
            ("RUNNING_MODERN.md", "新增现代运行说明、安装命令、验证命令和渲染器注意事项。"),
            ("decalib/models/FLAME.py", "为 chumpy/FLAME pickle 增加 inspect 与 NumPy 旧 API 兼容 shim。"),
            ("decalib/deca.py", "增加 PyTorch 2.6+ torch.load weights_only=False 兼容；支持 render_enabled=False。"),
            ("decalib/utils/renderer.py", "为 MPS 设备增加 PyTorch3D CPU rasterize 回退：模型可在 MPS，rasterizer 临时在 CPU。"),
            ("demos/demo_reconstruct.py", "新增 --device auto 与 --rendering False；无渲染模式自动禁用依赖渲染的输出项。"),
            ("decalib/datasets/datasets.py", "iscrop=False 时不再初始化 face-alignment 检测器，避免无必要下载/卡住。"),
            ("decalib/utils/util.py", "将 np.int 替换为 int，兼容 NumPy 2.x。"),
            ("decalib/trainer.py / decalib/utils/trainer.py", "训练加载 checkpoint 时兼容新版 torch.load。"),
        ],
        [2600, 6760],
    )

    add_heading(doc, "如何运行", 1)
    add_heading(doc, "1. 激活环境", 2)
    add_code_block(
        doc,
        [
            "cd /Users/houwingtsang/Documents/face_standardization_project/DECA",
            "conda activate deca-modern",
        ],
    )

    add_heading(doc, "2. 推荐的 Mac 无渲染运行方式", 2)
    add_para(doc, "适用于当前机器。它会跑通图像读取、模型推理、FLAME 几何与参数输出，但不会生成渲染可视化图、深度图或 OBJ。")
    add_code_block(
        doc,
        [
            "python demos/demo_reconstruct.py \\",
            "  -i TestSamples/examples \\",
            "  -s TestSamples/examples/results_modern_no_render \\",
            "  --device cpu \\",
            "  --rendering False \\",
            "  --iscrop False \\",
            "  --saveKpt True \\",
            "  --saveMat True",
        ],
    )

    add_heading(doc, "3. 单图验证命令", 2)
    add_code_block(
        doc,
        [
            "python demos/demo_reconstruct.py \\",
            "  -i TestSamples/examples/IMG_0392_inputs.jpg \\",
            "  -s TestSamples/examples/results_modern_no_render \\",
            "  --device cpu \\",
            "  --rendering False \\",
            "  --iscrop False \\",
            "  --saveKpt True \\",
            "  --saveMat True",
        ],
    )

    add_heading(doc, "4. 完整渲染输出路径", 2)
    add_para(doc, "如果要保存 saveDepth、saveObj、saveVis 或 saveImages，需要可用渲染器。本机已成功源码安装 PyTorch3D，因此 Mac 可以走 pytorch3d rasterizer。")
    add_heading(doc, "4.1 安装 PyTorch3D", 3)
    add_code_block(
        doc,
        [
            "conda activate deca-modern",
            "MACOSX_DEPLOYMENT_TARGET=10.14 CC=clang CXX=clang++ \\",
            "python -m pip install --no-build-isolation \"git+https://github.com/facebookresearch/pytorch3d.git\"",
        ],
    )
    add_para(doc, "普通 pip 源码安装曾因隔离构建环境看不到 torch 而失败；加 --no-build-isolation 后成功生成并安装 pytorch3d-0.7.9 的 macOS arm64 wheel。")
    add_heading(doc, "4.2 CPU + PyTorch3D 完整渲染", 3)
    add_code_block(
        doc,
        [
            "python demos/demo_reconstruct.py \\",
            "  -i TestSamples/examples/IMG_0392_inputs.jpg \\",
            "  -s TestSamples/examples/results_mac_rendered_obj \\",
            "  --device cpu \\",
            "  --rasterizer_type pytorch3d \\",
            "  --rendering True \\",
            "  --render_orig False \\",
            "  --iscrop False \\",
            "  --saveVis True \\",
            "  --saveDepth True \\",
            "  --saveObj True \\",
            "  --saveKpt True \\",
            "  --saveMat True",
        ],
    )
    add_heading(doc, "4.3 MPS 混合渲染", 3)
    add_para(doc, "直接让 PyTorch3D rasterizer 跑在 MPS 会报错：Cannot use CPU implementation: face_verts not on CPU。本次已加入兼容逻辑：模型主体使用 MPS，PyTorch3D rasterize 步骤临时回退 CPU。")
    add_code_block(
        doc,
        [
            "python demos/demo_reconstruct.py \\",
            "  -i TestSamples/examples/IMG_0392_inputs.jpg \\",
            "  -s TestSamples/examples/results_mac_rendered_mps \\",
            "  --device mps \\",
            "  --rasterizer_type pytorch3d \\",
            "  --rendering True \\",
            "  --render_orig False \\",
            "  --iscrop False \\",
            "  --saveVis True \\",
            "  --saveDepth True \\",
            "  --saveObj True \\",
            "  --saveKpt True \\",
            "  --saveMat True",
        ],
    )

    add_heading(doc, "验证结果", 1)
    add_para(doc, "已在 deca-modern 环境中完成三类单图 smoke test：无渲染、CPU 完整渲染、MPS 混合渲染。运行日志显示模型创建、权重加载、推理进度完成，并输出结果目录。")
    add_matrix(
        doc,
        ["检查项", "结果"],
        [
            ("import 检查", "torch、torchvision、cv2、skimage、scipy、yaml、face_alignment、kornia 均可导入。"),
            ("PyTorch3D 检查", "pytorch3d 可导入；版本 0.7.9；rasterize_meshes 入口可导入。"),
            ("硬件检查", "cuda False；mps True。"),
            ("模型加载", "成功加载 data/deca_model.tar。"),
            ("无渲染推理", "成功生成 IMG_0392_inputs_kpt2d.txt、IMG_0392_inputs_kpt3d.txt、IMG_0392_inputs.mat。"),
            ("CPU 完整渲染", "成功生成 vis.jpg、depth.jpg、obj、detail.obj、texture、normal map、关键点与 .mat。"),
            ("MPS 混合渲染", "成功生成完整渲染输出，但单图耗时约 5.18s，慢于 CPU 路线约 2.95s。"),
            ("提示信息", "torchvision pretrained 参数有弃用警告，但不阻断运行。"),
        ],
        [1900, 7460],
    )
    add_code_block(
        doc,
        [
            "TestSamples/examples/results_modern_no_render/IMG_0392_inputs/",
            "  IMG_0392_inputs_kpt2d.txt",
            "  IMG_0392_inputs_kpt3d.txt",
            "  IMG_0392_inputs.mat",
            "",
            "TestSamples/examples/results_mac_rendered_obj/IMG_0392_inputs/",
            "  IMG_0392_inputs.obj",
            "  IMG_0392_inputs_detail.obj",
            "  IMG_0392_inputs.mtl",
            "  IMG_0392_inputs.png",
            "  IMG_0392_inputs_normals.png",
            "  IMG_0392_inputs_depth.jpg",
            "  IMG_0392_inputs_kpt2d.txt",
            "  IMG_0392_inputs_kpt3d.txt",
            "  IMG_0392_inputs.mat",
        ],
    )

    add_heading(doc, "重要限制与建议", 1)
    for item in [
        "当前 Mac 已支持完整渲染，但推荐优先使用 --device cpu + --rasterizer_type pytorch3d，稳定且本次测试更快。",
        "MPS 可以跑，但本项目中 PyTorch3D rasterizer 仍需要 CPU 回退，因此整体不一定更快。",
        "如果输入图不是已经居中裁剪的人脸，--iscrop False 可能影响结果质量；开启 --iscrop True 会初始化 face-alignment 检测器，可能触发模型下载或较长初始化。",
        "standard rasterizer 仍依赖 CUDA，不适合 Mac；Linux + NVIDIA GPU 依然是原项目最正统的完整渲染路径。",
        "原始 requirements.txt 保留给 legacy Python 3.7/CUDA 10.1 路线；现代环境请优先看 RUNNING_MODERN.md。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "后续可选工作", 1)
    for step in [
        "将 PyTorch3D 安装命令和 MPS 混合渲染说明同步写入 RUNNING_MODERN.md。",
        "将 demo_transfer.py 与 demo_teaser.py 同步迁移到 --device auto / --rendering False / MPS fallback 风格。",
        "增加一个更小的 smoke_test.py，专门用于 CI 或本机快速检查模型加载。",
        "如果需要生产化使用，建议封装输入/输出目录、错误提示和依赖检查脚本。",
    ]:
        add_step(doc, step)

    add_heading(doc, "参考资料", 1)
    add_para(doc, "PyTorch 官方本地安装文档：https://docs.pytorch.org/get-started/locally/", size=10, color=MUTED)
    add_para(doc, "PyTorch3D 官方安装说明：https://github.com/facebookresearch/pytorch3d/blob/main/INSTALL.md", size=10, color=MUTED)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
