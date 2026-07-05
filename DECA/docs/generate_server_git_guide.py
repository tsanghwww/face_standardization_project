from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/houwingtsang/Documents/face_standardization_project/DECA/docs/服务器与Git仓库文件管理指引.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 45)
MUTED = RGBColor(91, 103, 112)
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"


def set_run(run, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def paragraph_border_bottom(paragraph, color="2E74B5", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text), size=10.5, color=INK)


def add_step(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text), size=10.5, color=INK)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_table_borders(table, color="D7DEE8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        set_run(p.add_run(line), size=9.3, color=RGBColor(35, 45, 55), font="Courier New")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(header), size=10, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(value), size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("DECA Server & Git Guide"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("Project collaboration guide"), size=9, color=MUTED)

    add_para(doc, "项目协作指引", size=10, color=MUTED, bold=True, after=2)
    add_para(doc, "DECA 服务器与 Git 仓库文件管理方案", size=22, color=INK, bold=True, after=4)
    add_para(doc, "适用范围：Mac 已验证环境、Python 3.11 + CUDA 兼容环境、本地服务器共享、Git 分支与文件归档规则", size=12, color=MUTED, after=10)
    meta = add_para(doc, f"建议版本日期：{date.today().isoformat()}  |  原则：代码进 Git，大文件进共享存储，渲染结果不重复计算", size=10, color=MUTED, after=10)
    paragraph_border_bottom(meta)

    doc.add_heading("1. 核心原则", level=1)
    add_para(doc, "本项目需要同时管理代码、模型资产、输入数据和大规模渲染结果。为了避免仓库膨胀、授权文件误传和团队重复 8 小时渲染，必须把文件分层管理。")
    for item in [
        "Git 仓库只放代码、配置、文档、轻量脚本和环境文件。",
        "DECA/FLAME 权重与授权资产放服务器 assets 目录，不直接提交 Git。",
        "输入数据集放服务器 datasets 目录，不直接提交 Git。",
        "10000 张图的渲染结果放服务器 results 目录，成员直接复用，不重复渲染。",
        "所有成员通过 manifest.csv 对齐 input、DECA 输出、gaze 输出和评估结果。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("2. 服务器目录建议", level=1)
    add_code_block(
        doc,
        [
            "/srv/face_standardization_project/",
            "├── repo/",
            "│   └── DECA/                         # Git 仓库：代码、配置、文档",
            "├── assets/",
            "│   └── deca_data/                    # 模型资产：不进 Git",
            "├── datasets/",
            "│   └── my_faces/                     # 输入图片：不进 Git",
            "├── results/",
            "│   └── my_faces_rendered/            # 渲染结果：不进 Git",
            "└── shared_packages/",
            "    ├── my_faces_rendered_light.tar.gz",
            "    └── my_faces_rendered_3d.tar.gz",
        ],
    )
    add_matrix(
        doc,
        ["目录", "用途", "是否进 Git"],
        [
            ("repo/DECA", "项目代码、环境文件、运行说明、工具脚本。", "是"),
            ("assets/deca_data", "deca_model.tar、generic_model.pkl、FLAME albedo、UV mask、模板等。", "否"),
            ("datasets/my_faces", "原始输入图片或数据集软链接。", "否"),
            ("results/my_faces_rendered", "批量渲染后的 vis/depth/obj/mat/关键点等结果。", "否"),
            ("shared_packages", "按 light/3d 分包后的共享压缩包。", "否"),
        ],
        [2300, 5260, 1800],
    )

    doc.add_heading("3. Git 仓库内部结构", level=1)
    add_code_block(
        doc,
        [
            "DECA/",
            "├── demos/",
            "├── decalib/",
            "├── configs/",
            "├── docs/",
            "├── scripts/",
            "├── envs/",
            "│   ├── environment_mac_py310.yml",
            "│   └── environment_cuda_py311.yml",
            "├── data/.gitkeep",
            "├── datasets/.gitkeep",
            "├── results/.gitkeep",
            "├── requirements_modern.txt",
            "├── RUNNING_MODERN.md",
            "└── README.md",
        ],
    )
    add_para(doc, "仓库内保留 data、datasets、results 的 .gitkeep，仅用于提示目录位置；真实大文件通过服务器目录或软链接接入。")

    doc.add_heading("4. 什么应该上传，什么不能上传", level=1)
    add_matrix(
        doc,
        ["类别", "应上传 Git", "不应上传 Git"],
        [
            ("代码", "demos/、decalib/、configs/、scripts/。", "运行中产生的临时文件、__pycache__、*.pyc。"),
            ("环境", "requirements_modern.txt、environment_modern.yml、envs/*.yml。", "本机 conda 环境目录、site-packages。"),
            ("文档", "README、RUNNING_MODERN、docs 下的项目说明。", "无需版本化的个人笔记。"),
            ("模型资产", "只上传下载说明和校验说明。", "data/*.tar、*.pkl、*.npz、模型权重、授权资产。"),
            ("数据与结果", "只上传 manifest 模板和字段说明。", "datasets/、results/、*.obj、*.mat、*_vis.jpg、*_depth.jpg。"),
        ],
        [1700, 3830, 3830],
    )

    doc.add_heading("5. 推荐 .gitignore", level=1)
    add_code_block(
        doc,
        [
            "# Model/data assets",
            "data/*",
            "!data/.gitkeep",
            "",
            "# Input datasets",
            "datasets/*",
            "!datasets/.gitkeep",
            "",
            "# Generated outputs",
            "results/*",
            "!results/.gitkeep",
            "logs/",
            "TestSamples/examples/results*/",
            "",
            "# Large generated files",
            "*.obj",
            "*.mtl",
            "*.mat",
            "*_depth.jpg",
            "*_vis.jpg",
            "*_normals.png",
            "",
            "# Python/cache",
            "__pycache__/",
            "*.pyc",
            ".DS_Store",
        ],
    )

    doc.add_heading("6. Python 3.11 + CUDA 兼容环境", level=1)
    add_para(doc, "另一台 CUDA 电脑应使用独立环境，不要复用 Mac 的 Python 3.10 环境。PyTorch3D 需要针对 Python 3.11 和 CUDA 机器重新安装或编译。")
    add_code_block(
        doc,
        [
            "conda create -n deca-cuda-py311 python=3.11 -y",
            "conda activate deca-cuda-py311",
            "pip install -r requirements_modern.txt",
            "pip install chumpy --no-build-isolation",
            "",
            "# 按 PyTorch 官方页面选择 CUDA wheel，例如：",
            "pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu126",
            "",
            "# PyTorch3D 需要在该环境中安装/编译：",
            "pip install --no-build-isolation \"git+https://github.com/facebookresearch/pytorch3d.git\"",
        ],
    )
    add_matrix(
        doc,
        ["环境", "推荐用途", "备注"],
        [
            ("Mac / Python 3.10", "当前已验证环境；适合本机 CPU + PyTorch3D 渲染。", "已验证无渲染、CPU 渲染、MPS 混合渲染。"),
            ("CUDA / Python 3.11", "另一台 GPU 电脑；适合加速推理和渲染。", "需要重新安装 PyTorch CUDA wheel 和 PyTorch3D。"),
            ("服务器共享目录", "团队复用数据、模型和渲染结果。", "不要让每个成员重复跑 8 小时。"),
        ],
        [2300, 3900, 3160],
    )

    doc.add_heading("7. 分支结构建议", level=1)
    add_matrix(
        doc,
        ["分支", "用途", "合并规则"],
        [
            ("main", "稳定主分支；放通用代码、文档、目录规范、环境文件。", "只合并验证过的改动。"),
            ("env/mac-py310-pytorch3d", "记录 Mac 已验证运行方案和兼容补丁。", "验证后合并回 main。"),
            ("env/cuda-py311", "测试 Python 3.11 + CUDA + PyTorch3D。", "CUDA 机器验证通过后合并。"),
            ("feature/batch-manifest", "生成 manifest、统计成功/失败、断点续跑脚本。", "功能稳定后合并。"),
            ("feature/evaluation-pipeline", "ArcFace、LPIPS、gaze consistency 等评估代码。", "由评估负责人维护。"),
            ("release/render-pipeline-v1", "阶段性稳定版本或交付版本。", "从 main 创建，不直接开发。"),
        ],
        [2600, 4360, 2400],
    )

    doc.add_heading("8. 成员如何接入服务器文件", level=1)
    add_para(doc, "成员 clone 仓库后，不应复制大文件到 Git 仓库内部；推荐用软链接接入服务器共享目录。")
    add_code_block(
        doc,
        [
            "git clone <repo-url> DECA",
            "cd DECA",
            "",
            "# 接入模型资产",
            "ln -s /srv/face_standardization_project/assets/deca_data data",
            "",
            "# 接入输入数据",
            "mkdir -p datasets results",
            "ln -s /srv/face_standardization_project/datasets/my_faces datasets/my_faces",
            "",
            "# 接入已渲染结果，避免重复 8 小时计算",
            "ln -s /srv/face_standardization_project/results/my_faces_rendered results/my_faces_rendered",
        ],
    )

    doc.add_heading("9. 避免重复渲染的协作流程", level=1)
    for step in [
        "由一台机器完成 DECA 批量渲染，输出到服务器 results/my_faces_rendered。",
        "渲染完成后统计 vis、depth、obj、detail.obj、mat 数量，确认接近输入图数量。",
        "生成 manifest.csv，把每张输入图和对应输出文件路径一一对应。",
        "团队成员直接读取 results 和 manifest，不重新跑 DECA。",
        "Person B 使用原图和 manifest 跑 gaze estimator；Person C 使用 manifest 和 outputs 做指标评估。",
    ]:
        add_step(doc, step)
    add_code_block(
        doc,
        [
            "find results/my_faces_rendered -name \"*_vis.jpg\" | wc -l",
            "find results/my_faces_rendered -name \"*_depth.jpg\" | wc -l",
            "find results/my_faces_rendered -name \"*.obj\" | wc -l",
            "find results/my_faces_rendered -name \"*_detail.obj\" | wc -l",
            "find results/my_faces_rendered -name \"*.mat\" | wc -l",
        ],
    )

    doc.add_heading("10. 推荐交付包", level=1)
    add_matrix(
        doc,
        ["交付包", "内容", "适用成员"],
        [
            ("results_light", "vis.jpg、depth.jpg、kpt2d.txt、kpt3d.txt、mat。", "评估、写作、快速检查。"),
            ("results_3d", "obj、detail.obj、mtl、texture png、normals png。", "需要打开 3D 模型或做几何分析的成员。"),
            ("manifest.csv", "image_id、input_path、vis_path、depth_path、normal_path、obj_path、mat_path、status。", "所有成员。"),
            ("assets package", "DECA/FLAME 必要资产或内部下载链接。", "需要复现运行的成员。"),
        ],
        [1900, 5160, 2300],
    )

    doc.add_heading("11. 关键结论", level=1)
    add_para(doc, "项目协作的核心不是让所有人都重新跑 DECA，而是一次性完成批量渲染后，把结果作为共享中间产物。Git 管代码，服务器管大文件，manifest 管路径索引。这样 Person A 的 8 小时渲染不会被团队重复浪费，Person B 和 Person C 可以直接进入 gaze pipeline、evaluation 和 report 工作。", bold=True)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
